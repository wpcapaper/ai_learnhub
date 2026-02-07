"""
DOCX转JSON转换脚本
从Word文档中提取题目，识别红色标记的答案，转换为JSON格式供导入
"""
import sys
import re
import json
import argparse
from pathlib import Path
from typing import List, Dict, Optional, Tuple
from docx import Document
from docx.shared import RGBColor

class DocxQuestionParser:
    """DOCX题目解析器"""

    def __init__(self, docx_path: str):
        """
        初始化解析器

        Args:
            docx_path: DOCX文件路径
        """
        self.doc = Document(docx_path)
        self.docx_path = Path(docx_path)
        self.questions = []
        self.current_section = None
        self.current_question = None
        self.current_options = []

    def parse(self) -> List[Dict]:
        """
        解析文档并返回题目列表

        Returns:
            List[Dict]: 题目列表
        """
        i = 0
        while i < len(self.doc.paragraphs):
            para = self.doc.paragraphs[i]
            text = para.text.strip()

            # 跳过空行
            if not text:
                i += 1
                continue

            # 检测章节标题（如"一、单选题"）
            if self._is_section_header(text):
                # 保存上一道题（如果有）
                self._save_current_question()
                self.current_section = text
                self.current_question = None
                self.current_options = []
                i += 1
                continue

            # 检测题目（如"1、题目内容"）
            if self._is_question(text):
                # 保存上一道题（如果有）
                self._save_current_question()

                # 开始新题目
                self.current_question = {
                    'number': self._extract_question_number(text),
                    'content': self._extract_question_text(text),
                    'type': self._determine_question_type()
                }
                self.current_options = []

                i += 1
                continue

            # 判断题特殊处理：检查是否为"正确错误"格式
            if (self.current_question and
                self.current_question['type'] == 'true_false' and
                '正确错误' in text):

                # 检查红色文本标记
                correct_answer = self._extract_true_false_answer(para)
                if correct_answer:
                    # 根据schema要求，判断题必须提供options字段
                    # 格式: {"A": "对", "B": "错"}
                    # correct_answer: "A"（对）或 "B"（错）
                    if correct_answer == '对':
                        options_dict = {"A": "对", "B": "错"}
                        answer_letter = "A"
                    else:  # correct_answer == '错'
                        options_dict = {"A": "对", "B": "错"}
                        answer_letter = "B"

                    # 直接创建判断题
                    self.questions.append({
                        'question_type': 'true_false',
                        'content': self.current_question['content'],
                        'options': options_dict,
                        'correct_answer': answer_letter,
                        'explanation': '',
                        'difficulty': 2,
                        'knowledge_points': [],
                        'metadata': {
                            'source': 'docx',
                            'docx_file': self.docx_path.name
                        }
                    })

                    # 清空当前题目
                    self.current_question = None
                    self.current_options = []

                i += 1
                continue

            # 检测选项（如"A. 选项内容"）
            if self._is_option(text) and self.current_question:
                option_data = self._extract_option(para)
                if option_data:
                    self.current_options.append(option_data)

                i += 1
                continue

            # 其他内容，跳过
            i += 1

        # 保存最后一道题（如果有）
        self._save_current_question()

        return self.questions

    def _is_section_header(self, text: str) -> bool:
        """
        判断是否为章节标题

        Args:
            text: 段落文本

        Returns:
            bool: 是否为章节标题
        """
        return bool(re.match(r'^[一二三四五六七八九十]+、\w+', text))

    def _is_question(self, text: str) -> bool:
        """
        判断是否为题目

        Args:
            text: 段落文本

        Returns:
            bool: 是否为题目
        """
        return bool(re.match(r'^\d+、', text))

    def _is_option(self, text: str) -> bool:
        """
        判断是否为选项

        Args:
            text: 段落文本

        Returns:
            bool: 是否为选项
        """
        # 匹配 "A. 选项" 或 " A. 选项"（前面可能有空格）
        return bool(re.match(r'^\s*[A-Z]\.', text))

    def _extract_question_number(self, text: str) -> int:
        """
        提取题目编号

        Args:
            text: 题目文本

        Returns:
            int: 题目编号
        """
        match = re.match(r'^(\d+)、', text)
        return int(match.group(1)) if match else 0

    def _extract_question_text(self, text: str) -> str:
        """
        提取题目内容（去掉编号）

        Args:
            text: 题目文本

        Returns:
            str: 题目内容
        """
        return re.sub(r'^\d+、', '', text).strip()

    def _extract_option(self, paragraph) -> Optional[Dict]:
        """
        提取选项内容并检测是否为红色

        Args:
            paragraph: docx段落对象

        Returns:
            Optional[Dict]: 选项数据，格式为 {'letter': 'A', 'text': '内容', 'is_red': bool}
        """
        text = paragraph.text.strip()

        # 提取选项字母
        match = re.match(r'^\s*([A-Z])\.\s*', text)
        if not match:
            return None

        letter = match.group(1)
        option_text = re.sub(r'^\s*[A-Z]\.\s*', '', text).strip()

        # 检测是否为红色
        is_red = self._is_red_text(paragraph)

        return {
            'letter': letter,
            'text': option_text,
            'is_red': is_red
        }

    def _extract_true_false_answer(self, paragraph) -> Optional[str]:
        """
        从"正确错误"段落中提取正确答案（基于红色标记）

        Args:
            paragraph: docx段落对象

        Returns:
            Optional[str]: '对' 或 '错'
        """
        for run in paragraph.runs:
            if run.font.color and run.font.color.rgb:
                if run.font.color.rgb == RGBColor(0xFF, 0x00, 0x00):
                    text = run.text.strip()
                    # 判断是"对"还是"错"
                    if '对' in text or '正确' in text or '是' in text:
                        return '对'
                    elif '错' in text or '错误' in text or '否' in text:
                        return '错'
        return None

    def _is_red_text(self, paragraph) -> bool:
        """
        检测段落是否包含红色文本

        Args:
            paragraph: docx段落对象

        Returns:
            bool: 是否为红色
        """
        for run in paragraph.runs:
            if run.font.color and run.font.color.rgb:
                if run.font.color.rgb == RGBColor(0xFF, 0x00, 0x00):
                    return True
        return False

    def _determine_question_type(self) -> str:
        """
        根据章节标题确定题目类型

        Returns:
            str: 题目类型（single_choice, multiple_choice, true_false）
        """
        if not self.current_section:
            return 'single_choice'

        section_text = self.current_section

        if '单选' in section_text:
            return 'single_choice'
        elif '多选' in section_text:
            return 'multiple_choice'
        elif '判断' in section_text:
            return 'true_false'
        else:
            return 'single_choice'

    def _save_current_question(self):
        """保存当前题目到列表"""
        if not self.current_question:
            return

        # 根据题目类型处理
        question_type = self.current_question['type']

        if question_type == 'single_choice':
            # 单选题：正确答案是红色的选项
            correct_answer = None
            for opt in self.current_options:
                if opt['is_red']:
                    correct_answer = opt['letter']
                    break

            if not correct_answer:
                print(f"⚠️  警告：第{self.current_question['number']}题未找到正确答案（红色选项）")
                return

            # 构建选项字典
            options_dict = {opt['letter']: opt['text'] for opt in self.current_options}

            # 添加到结果
            self.questions.append({
                'question_type': 'single_choice',
                'content': self.current_question['content'],
                'options': options_dict,
                'correct_answer': correct_answer,
                'explanation': '',
                'difficulty': 2,
                'knowledge_points': [],
                'metadata': {
                    'source': 'docx',
                    'docx_file': self.docx_path.name
                }
            })

        elif question_type == 'multiple_choice':
            # 多选题：正确答案是多个红色的选项
            correct_answers = []
            for opt in self.current_options:
                if opt['is_red']:
                    correct_answers.append(opt['letter'])

            if not correct_answers:
                print(f"⚠️  警告：第{self.current_question['number']}题未找到正确答案（红色选项）")
                return

            # 构建选项字典
            options_dict = {opt['letter']: opt['text'] for opt in self.current_options}

            # ✅ 按字母顺序拼接答案
            sorted_answers = sorted(correct_answers)
            
            # ✅ 使用逗号分隔的格式（符合schema要求）
            # 例如: "AB,CD" 或 "A,C,D"
            # 也可以使用数组格式: ["A", "B", "C", "D"]
            # 这里选择字符串格式，用逗号分隔
            correct_answer_str = ','.join(sorted_answers)

            # 添加到结果
            self.questions.append({
                'question_type': 'multiple_choice',
                'content': self.current_question['content'],
                'options': options_dict,
                'correct_answer': correct_answer_str,
                'explanation': '',
                'difficulty': 2,
                'knowledge_points': [],
                'metadata': {
                    'source': 'docx',
                    'docx_file': self.docx_path.name
                }
            })

        elif question_type == 'true_false':
            # 判断题：需要检查是否有"对"/"错"的红色标记
            # 根据分析，判断题可能是选项形式的（A. 对 B. 错）
            # 或者是直接标记为对/错的文本

            # 先尝试选项方式（如果当前有选项）
            if self.current_options:
                correct_answer = None
                for opt in self.current_options:
                    if opt['is_red']:
                        # 判断是"对"还是"错"
                        opt_text = opt['text'].strip()
                        if '对' in opt_text or '正确' in opt_text or '是' in opt_text:
                            correct_answer = '对'
                        elif '错' in opt_text or '错误' in opt_text or '否' in opt_text:
                            correct_answer = '错'
                        break

                if correct_answer:
                    options_dict = {opt['letter']: opt['text'] for opt in self.current_options}

                    self.questions.append({
                        'question_type': 'true_false',
                        'content': self.current_question['content'],
                        'options': options_dict,
                        'correct_answer': correct_answer,
                        'explanation': '',
                        'difficulty': 2,
                        'knowledge_points': [],
                'metadata': {
                    'source': 'docx',
                    'docx_file': self.docx_path.name
                }
                    })
                    return

            # 如果没有找到正确答案，跳过该题
            print(f"⚠️  警告：第{self.current_question['number']}判断题未找到正确答案")


def main():
    """主函数"""
    parser = argparse.ArgumentParser(
        description='将 DOCX 文件转换为 JSON 格式的题目数据',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog='''
示例:
  # 转换单个文件（使用默认输入和输出目录）
  python convert_docx_to_json.py -f questions.docx

  # 指定输入目录
  python convert_docx_to_json.py -f questions.docx -i custom/input

  # 指定输出目录
  python convert_docx_to_json.py -f questions.docx -o custom/output

  # 转换并指定占位符
  python convert_docx_to_json.py -f questions.docx -p "暂无解析"
        '''
    )

    parser.add_argument(
        '-f', '--file',
        type=str,
        default='questions.docx',
        help='输入文件名（默认: questions.docx）。文件应位于 scripts/data/input/ 目录'
    )

    parser.add_argument(
        '-i', '--input-dir',
        type=str,
        default=None,
        help='输入目录路径（默认: scripts/data/input/）'
    )

    parser.add_argument(
        '-o', '--output-dir',
        type=str,
        default=None,
        help='输出目录路径（默认: scripts/data/output/）'
    )

    parser.add_argument(
        '-p', '--placeholder-explanation',
        default='暂无解析',
        help='解析字段的占位符文本（默认：暂无解析）'
    )

    parser.add_argument(
        '-d', '--default-difficulty',
        type=int,
        default=2,
        help='默认难度等级（1-5，默认：2）'
    )

    args = parser.parse_args()

    # 设置路径
    script_dir = Path(__file__).parent
    input_dir = Path(args.input_dir) if args.input_dir else script_dir / "data" / "input"
    output_dir = Path(args.output_dir) if args.output_dir else script_dir / "data" / "output"

    print(f"脚本目录: {script_dir}")
    print(f"输入目录: {input_dir}")
    print(f"输出目录: {output_dir}\n")

    # 构建输入文件路径
    docx_path = input_dir / args.file

    if not docx_path.exists():
        print(f"❌ 文件不存在: {docx_path}")
        print(f"\n提示：请确保文件位于 {input_dir} 目录下")
        sys.exit(1)

    # 设置输出路径
    output_path = output_dir / f"{docx_path.stem}.json"

    # 创建解析器并解析
    print(f"📖 正在解析: {docx_path}")
    parser = DocxQuestionParser(docx_path)
    questions = parser.parse()

    if not questions:
        print("❌ 错误：未找到任何题目")
        sys.exit(1)

    # 更新占位符和难度
    for q in questions:
        q['explanation'] = args.placeholder_explanation
        q['difficulty'] = args.default_difficulty

    # 统计信息
    type_counts = {}
    for q in questions:
        qtype = q['question_type']
        type_counts[qtype] = type_counts.get(qtype, 0) + 1

    print(f"\n✅ 解析完成!")
    print(f"  总题目数: {len(questions)}")
    print(f"  单选题: {type_counts.get('single_choice', 0)}")
    print(f"  多选题: {type_counts.get('multiple_choice', 0)}")
    print(f"  判断题: {type_counts.get('true_false', 0)}")

    # 确保输出目录存在
    output_path.parent.mkdir(parents=True, exist_ok=True)

    # 写入JSON文件
    with open(output_path, 'w', encoding='utf-8') as f:
        json.dump(questions, f, ensure_ascii=False, indent=2)

    print(f"\n📄 已保存到: {output_path}")

    # 验证文件
    try:
        with open(output_path, 'r', encoding='utf-8') as f:
            json.load(f)
        print("✅ JSON文件验证通过")
    except json.JSONDecodeError as e:
        print(f"❌ JSON文件验证失败: {e}")
        sys.exit(1)


if __name__ == "__main__":
    main()
